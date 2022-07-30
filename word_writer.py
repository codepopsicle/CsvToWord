from docx import Document
from docx.shared import Cm, Pt

word_document = Document()
# Add a Title to the document
word_document.add_heading('Addresses', 0)
document_name = 'Addresses_Doc'


def writeToRow(word_table, rowData):
    row = word_table.add_row().cells
    row[0].text = str(rowData['first_name'])
    row[1].text = str(rowData['last_name'])
    row[2].text = str(rowData['address'])
    row[3].text = str(rowData['street'])
    row[4].text = str(rowData['state'])
    row[5].text = str(rowData['zip'])
    word_document.add_page_break()


def createTable():
    # creating the table
    table = word_document.add_table(0, 0)  # we add rows iteratively
    table.style = 'Colorful List'
    first_column_width = 5
    second_column_with = 10
    third_column_width = 5
    fourth_column_with = 10
    fifth_column_width = 5
    sixth_column_width = 10
    table.add_column(Cm(first_column_width))
    table.add_column(Cm(second_column_with))
    table.add_column(Cm(third_column_width))
    table.add_column(Cm(fourth_column_with))
    table.add_column(Cm(fifth_column_width))
    table.add_column(Cm(sixth_column_width))

    # Adding heading in the 1st row of the table
    row = table.add_row().cells
    row[0].text = 'First Name'
    row[1].text = 'Last Name'
    row[2].text = 'Address'
    row[3].text = 'Street'
    row[4].text = 'State'
    row[5].text = 'Zip'

    return table

def saveWordFile():
    word_document.save(document_name + '.docx')
